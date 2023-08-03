import argparse
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.probability import FreqDist
from sklearn.feature_extraction.text import TfidfVectorizer
import re
import sys
import fitz
from collections import defaultdict
from bs4 import BeautifulSoup
import docx
import os

# Télécharger les données nltk sans afficher les messages
nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)

def exclure_sommaire(texte):
    # Motifs pour détecter le sommaire
    motifs_sommaire = [
        r"Sommaire\s*",
        r"Table\s+des\s+matières\s*",
        r"Table\s+of\s+Contents\s*",
        r"Liste\s+des\s+figures\s*",
        r"Liste\s+des\s+tableaux\s*",
        r"Chapitre\d+\s*:\s*.*\d+",
        r"Chapitre\d+\s*:\s*.*?\d+",
        r"Chapitre\d+\s*:\s*.*?(?=\d)",
        r"Chapitre\s+\d+\s*:\s*.*?(?=\d)",
        r"Conclusion Générale\s*.*?\d+",
        r"Figure\s+\d+:\s*.*?(?=\d)",
        r"Tableau\s+\d+:\s*.*?(?=\d)",
        r"Bibliographie et Néographie\s*",
        r"\[\d+\]\s+https?://[^\s]+",
        r"ANNEXE\d+\s*:\s*.*",
        r"ANNEXE\s+\d+\s*:\s*.*",
        r"Remerciements\s*",
        r"Dédicaces\s*"
    ]
    # Utiliser le mode re.DOTALL pour capturer les motifs multi-lignes
    sommaire_match = None
    for motif in motifs_sommaire:
        sommaire_match = re.search(motif, texte, re.IGNORECASE | re.DOTALL)
        if sommaire_match:
            break

    if sommaire_match:
        debut_sommaire = sommaire_match.start()
        fin_sommaire = sommaire_match.end()
        return texte[:debut_sommaire] + texte[fin_sommaire:]
    
    return texte

def process_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()
    sentences, words = preprocess_text(text)
    word_frequencies = calculate_word_frequency(words)
    summary = generate_summary(sentences, word_frequencies)
    return summary

def process_pdf(file_path, num_sentences=3, short_text_threshold=10):
    with fitz.open(file_path) as pdf_document:
        text = ""
        for page_number in range(pdf_document.page_count):
            page = pdf_document[page_number]
            page_text = page.get_text()
            if exclure_sommaire(page_text) != page_text:
                continue
            text += page_text

    # Tokenize the text into sentences
    sentences = sent_tokenize(text)

    # Get the total number of sentences in the text
    total_sentences = len(sentences)

    # Check if the text is short
    if total_sentences <= short_text_threshold:
        num_sentences = min(total_sentences, num_sentences)  # Use the actual number of sentences if fewer than num_sentences
    else:
        num_sentences = 3  # Use the default num_sentences for longer texts

    # Preprocess text to perform TF-IDF vectorization
    tfidf_vectorizer = TfidfVectorizer(stop_words="english")
    tfidf_matrix = tfidf_vectorizer.fit_transform(sentences)
    word_frequencies = dict(zip(tfidf_vectorizer.get_feature_names_out(), tfidf_matrix.sum(axis=0).tolist()[0]))

    # Generate the summary
    summary = generate_summary(sentences, word_frequencies, num_sentences)

    return summary

def process_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    sentences, words = preprocess_text(text)
    word_frequencies = calculate_word_frequency(words)
    summary = generate_summary(sentences, word_frequencies)
    return summary

def preprocess_text(text):
    sentences = sent_tokenize(text)
    words = word_tokenize(text.lower())
    stop_words = set(stopwords.words("english"))
    filtered_words = [word for word in words if word.isalnum() and word not in stop_words]
    return sentences, filtered_words

def calculate_word_frequency(words):
    fdist = FreqDist(words)
    return fdist

def generate_summary(sentences, word_frequencies, num_sentences=3):
    ranked_sentences = [(sentences[i], sum(word_frequencies.get(word.lower(), 0) for word in word_tokenize(sentences[i].lower())))
                        for i in range(len(sentences))]
    ranked_sentences = sorted(ranked_sentences, key=lambda x: x[1], reverse=True)
    summary = [ranked_sentences[i][0] for i in range(min(num_sentences, len(ranked_sentences)))]
    return ' '.join(summary)

def titles_from_text(text):
    titles = defaultdict(list)

    # Add your regex patterns for chapter, figure, and table titles here
    pattern1 = r"Chapitre(\d+)\s*:\s*(.*?)\s*(\d+)"
    pattern2 = r"Chapitre\s*(.*?)\s*:\s*(.*)"
    pattern3 = r"\d+\s*:\s*(.*)"
    pattern4 = r"Figure\s+\d+\s*:\s*(.*?)\s*"
    pattern5 = r"Tableau\s+\d+\s*:\s*(.*?)\s*"

    # Find matches for each pattern and store the titles in the 'titles' dictionary
    matches1 = re.finditer(pattern1, text, re.IGNORECASE)
    for match in matches1:
        chapter_number = match.group(1)
        chapter_title = match.group(3)
        titles[chapter_number].append(chapter_title)

    matches2 = re.finditer(pattern2, text, re.IGNORECASE)
    for match in matches2:
        chapter_title = match.group(2)
        titles["Chapters"].append(chapter_title)

    matches3 = re.finditer(pattern3, text, re.IGNORECASE)
    for match in matches3:
        chapter_title = match.group(1)
        titles["Chapters"].append(chapter_title)

    matches4 = re.finditer(pattern4, text, re.IGNORECASE)
    for match in matches4:
        figure_title = match.group(1)
        titles["Figures"].append(figure_title)

    matches5 = re.finditer(pattern5, text, re.IGNORECASE)
    for match in matches5:
        table_title = match.group(1)
        titles["Tables"].append(table_title)

    return titles

def process_file(file_path, num_sentences=3, short_text_threshold=10):
    _, file_extension = os.path.splitext(file_path)

    if file_extension == '.pdf':
        return process_pdf(file_path, num_sentences, short_text_threshold)
    elif file_extension == '.docx':
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Use the 'titles_from_text' function to get the titles from the text
        titles = titles_from_text(text)
        # Convert titles dictionary to a string to include in the summary
        titles_str = "\n".join(f"{key}: {', '.join(value)}" for key, value in titles.items())

        sentences, words = preprocess_text(text)
        word_frequencies = calculate_word_frequency(words)
        summary = generate_summary(sentences, word_frequencies, num_sentences)

        # Add the titles_str to the beginning of the summary
        final_summary = titles_str + "\n\n" + summary

        return final_summary
    elif file_extension == '.txt':
        return process_txt(file_path)
    else:
        print("Unsupported file format. Only PDF, TXT, and DOCX files are supported.")
        sys.exit(1)

def main():
    parser = argparse.ArgumentParser(description="Generate summary from a text file, a Word (DOCX) file, or a PDF file.")
    parser.add_argument("file_path", type=str, help="Path to the input file.")
    args = parser.parse_args()

    if args.file_path.endswith(('.pdf', '.txt', '.docx')):
        summary = process_file(args.file_path)
    else:
        print("Unsupported file format. Only PDF, TXT, and DOCX files are supported.")
        sys.exit(1)

    # Utiliser sys.stdout.buffer.write pour afficher correctement les caractères spéciaux
    sys.stdout.buffer.write(summary.encode('utf-8'))

if __name__ == "__main__":
    main()